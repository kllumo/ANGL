import re
import pandas as pd
import sys
import os
import spacy
import docx
import fitz  # PyMuPDF
from PIL import Image, ImageFilter

nlp = spacy.load("ru_core_news_md")

# ---------------------- Маскировка Текстов ----------------------
def mask_kazakh_names(text):
    text = re.sub(r'(Аты:\s*)(\w+)', r'\1Ф.И.О.', text, flags=re.IGNORECASE)
    text = re.sub(r'(Тегі:\s*)(\w+)', r'\1Ф.И.О.', text, flags=re.IGNORECASE)
    text = re.sub(r'(Фамилия:\s*)(\w+)', r'\1Ф.И.О.', text, flags=re.IGNORECASE)
    text = re.sub(r'(Есімі:\s*)(\w+)', r'\1Ф.И.О.', text, flags=re.IGNORECASE)
    return text

def mask_entities_with_spacy(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PER":
            text = text.replace(ent.text, "Ф.И.О.")
        elif ent.label_ in ["LOC", "GPE"]:
            text = text.replace(ent.text, "[MASKED ADDRESS]")
        elif ent.label_ == "ORG":
            text = text.replace(ent.text, "[MASKED ORG]")
    return text

def mask_sensitive_text(text):
    text = mask_kazakh_names(text)
    text = mask_entities_with_spacy(text)
    patterns = {
        r'\b\d{12}\b': '000000000000',
        r'\bKZ\d{2}(?:\s?\d{4}){4}\b': 'KZ000000000000000000',
        r'\b(?:\d{4}\s){3}\d{4}\b': '0000 0000 0000 0000',
        r'\b\d{10,20}\b': '00000000000000000000',
        r'\b\d{7,9}\b': '000000000',
        r'\b\S+@\S+\.\S+\b': 'hidden@example.com',
        r'\b\d{2}[./-]\d{2}[./-]\d{4}\b': '00.00.0000',
        r'\b\d{2}[./-]\d{2}[./-]\d{4}\s*[-–]\s*\d{2}[./-]\d{2}[./-]\d{4}\b': '00.00.0000 - 00.00.0000'
    }
    for pattern, repl in patterns.items():
        text = re.sub(pattern, repl, text)
    text = re.sub(r'\b[A-Z][a-z]+ [A-Z][a-z]+(?: [A-Z][a-z]+)?\b', 'NAME', text)
    address_markers = r'(ул\.|улица|г\.|город|адрес|street|avenue|road|проспект|пер\.|переулок)'
    text = re.sub(f'{address_markers}.*', r'\1 [MASKED ADDRESS]', text, flags=re.IGNORECASE)
    return text

def anonymize_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    masked_content = mask_sensitive_text(content)
    out_path = f"{os.path.splitext(file_path)[0]}_anonymized.txt"
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(masked_content)
    return out_path
    #print(f"✅ TXT сохранён: {out_path}")

def anonymize_csv_file(file_path):
    df = pd.read_csv(file_path)
    for col in df.columns:
        if col.strip().lower() in ['фамилия', 'имя', 'отчество']:
            df[col] = 'Ф.И.О.'
        else:
            df[col] = df[col].astype(str).apply(mask_sensitive_text)
    out_path = f"{os.path.splitext(file_path)[0]}_anonymized.csv"
    df.to_csv(out_path, index=False)
    return out_path
    #print(f"✅ CSV сохранён: {out_path}")

def anonymize_excel_file(file_path):
    df = pd.read_excel(file_path)
    for col in df.columns:
        if col.strip().lower() in ['фамилия', 'имя', 'отчество']:
            df[col] = 'Ф.И.О.'
        else:
            df[col] = df[col].astype(str).apply(mask_sensitive_text)
    out_path = f"{os.path.splitext(file_path)[0]}_anonymized.xlsx"
    df.to_excel(out_path, index=False)
    return out_path
    #print(f"✅ Excel сохранён: {out_path}")

def anonymize_docx_file(file_path):
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        para.text = mask_sensitive_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = mask_sensitive_text(cell.text)
    out_path = f"{os.path.splitext(file_path)[0]}_anonymized.docx"
    doc.save(out_path)
    return out_path
    #print(f"✅ DOCX сохранён: {out_path}")

# ---------------- Маскировка PDF с координатами и блюром ----------------
MASKS_ORIGINAL = [
    (238.5, 162.56, 297.08, 181.88, 'Ф.И.О.'),
    (238.5, 207.56, 290.88, 226.88, ''),
    (238.5, 252.56, 330.88, 271.88, ''),
    (238.5, 297.56, 317.98, 316.88, '00.00.0000'),
    (111.0, 342.59, 228.73, 364.31, '000000000000'),
    (396.0, 402.96, 477.0, 426.97, '000000000'),
    (186.0, 447.56, 250.30, 466.88, '[MASKED]'),
    (186.0, 481.31, 233.97, 500.63, '[MASKED]'),
    (186.0, 515.81, 402.55, 535.13, '[MASKED]'),
    (186.0, 550.31, 360.76, 569.63, '00.00.0000 - 00.00.0000'),
    (73.5, 586.34, 518.43, 660.41, '[MACHINE ZONE MASKED]')  # Машинная зона (оставляем без изменений)
]

# Увеличиваем ширину всех кроме последнего на 40%
MASKS = []
for idx, (x0, y0, x1, y1, replacement) in enumerate(MASKS_ORIGINAL):
    if idx == len(MASKS_ORIGINAL) - 1:  # Последний элемент (машинная зона)
        MASKS.append((x0, y0, x1, y1, replacement))
    else:
        new_x1 = x0 + (x1 - x0) * 2
        MASKS.append((x0, y0, new_x1, y1, replacement))
PHOTO_BLUR_RECT = (60, 135, 210, 325)

def apply_blur_to_area(page, blur_rect):
    pix = page.get_pixmap()
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    zoom_x = pix.width / page.rect.width
    zoom_y = pix.height / page.rect.height
    x0, y0, x1, y1 = blur_rect
    left, top = int(x0 * zoom_x), int(y0 * zoom_y)
    right, bottom = int(x1 * zoom_x), int(y1 * zoom_y)
    region = img.crop((left, top, right, bottom))
    region = region.filter(ImageFilter.GaussianBlur(radius=15))
    img.paste(region, (left, top))
    temp_image_path = "temp_blurred_page.png"
    img.save(temp_image_path)
    new_doc = fitz.open()
    img_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
    img_page.insert_image(page.rect, filename=temp_image_path)
    return new_doc

def anonymize_pdf_file(file_path):
    doc = fitz.open(file_path)
    page = doc[0]
    for x0, y0, x1, y1, replacement in MASKS:
        rect = fitz.Rect(x0, y0, x1, y1)
        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
        if replacement:
            page.insert_textbox(rect, replacement, fontsize=10, color=(0, 0, 0), align=1)
    temp_path = f"{os.path.splitext(file_path)[0]}_temp.pdf"
    doc.save(temp_path)
    doc.close()
    temp_doc = fitz.open(temp_path)
    blurred_doc = apply_blur_to_area(temp_doc[0], PHOTO_BLUR_RECT)
    out_path = f"{os.path.splitext(file_path)[0]}_final_anonymized.pdf"
    blurred_doc.save(out_path)
    blurred_doc.close()
    os.remove(temp_path)
    return out_path
    #print(f"✅ PDF сохранён: {out_path}")

# ------------------------- MAIN -------------------------
def anonymize_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        return anonymize_text_file(file_path)
    elif ext == '.csv':
        return anonymize_csv_file(file_path)
    elif ext == '.xlsx':
        return anonymize_excel_file(file_path)
    elif ext == '.docx':
        return anonymize_docx_file(file_path)
    elif ext == '.pdf':
        return anonymize_pdf_file(file_path)
    else:
        print("❌ Поддерживаются только .txt, .csv, .xlsx, .docx и .pdf")
